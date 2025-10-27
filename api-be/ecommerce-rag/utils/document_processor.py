import os
from pathlib import Path
from typing import List
from docx import Document
from langchain.schema import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config.config import Config

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
        )
    
    def load_docx(self, file_path: str) -> str:
        """Đọc file Word và trích xuất text"""
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Đọc text từ tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    
    def process_document(self, file_path: str, metadata: dict = None) -> List[LangchainDocument]:
        """Xử lý document và chia thành chunks"""
        # Đọc nội dung file
        if file_path.endswith('.docx'):
            content = self.load_docx(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        # Tạo metadata
        if metadata is None:
            metadata = {}
        metadata['source'] = file_path
        metadata['filename'] = os.path.basename(file_path)
        
        # Tạo document
        doc = LangchainDocument(page_content=content, metadata=metadata)
        
        # Chia thành chunks
        chunks = self.text_splitter.split_documents([doc])
        
        return chunks